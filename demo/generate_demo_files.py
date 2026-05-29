"""Generate demo vendor-quote attachments in every supported upload format.

These are the files you hand-upload during a live demo to show AerOS's
format-agnostic intake. Every file quotes the SAME four line items from
RFx #1 ("Weekly Dairy & Produce Replenishment - W23") so the offers fuse
cleanly into the side-by-side comparison matrix:

    Tomato            200 kg     (target 18.00)
    Onion             150 kg     (target 22.00)
    Full Cream Milk   300 ltr    (target 56.00)
    Paneer             50 kg     (target 320.00)

Prices differ per vendor on purpose so "lowest price" and "best lead time"
highlights are visible. Run:

    .venv/bin/python demo/generate_demo_files.py

Output lands in demo/attachments/. Supported MIME types covered:
PDF, XLSX, DOCX, CSV, TXT, HTML, PNG, JPEG, WEBP.
"""

import csv
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "attachments")
os.makedirs(OUT, exist_ok=True)

ITEMS = ["Tomato", "Onion", "Full Cream Milk", "Paneer"]
UNITS = ["kg", "kg", "ltr", "kg"]
QTY = [200, 150, 300, 50]


def _rows(prices):
    """Zip items/qty/unit/price into rows with computed totals."""
    return [
        (ITEMS[i], QTY[i], UNITS[i], prices[i], round(QTY[i] * prices[i], 2))
        for i in range(len(ITEMS))
    ]


def csv_freshfarm():
    """FreshFarm Dairy — cheapest milk/paneer, doorstep."""
    path = os.path.join(OUT, "freshfarm_quote.csv")
    with open(path, "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["FreshFarm Dairy Quotation for RFx W23"])
        w.writerow(["Item", "Qty", "Unit", "Unit Price (INR)", "Total (INR)"])
        for r in _rows([16.50, 20.00, 52.00, 290.00]):
            w.writerow(r)
        w.writerow(["", "", "", "Grand Total", 36400.00])
        w.writerow(["Payment Terms", "NET15", "", "", ""])
        w.writerow(["Delivery", "Doorstep, within 18 hours", "", "", ""])
    return path


def xlsx_sabzimandi():
    """Sabzi Mandi Co — sharpest produce prices."""
    from openpyxl import Workbook

    path = os.path.join(OUT, "sabzimandi_quote.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Quotation"
    ws.append(["Sabzi Mandi Co — Fresh Produce Quote"])
    ws.append(["Item", "Qty", "Unit", "Unit Price (INR)", "Total (INR)"])
    for r in _rows([15.00, 18.50, 55.00, 305.00]):
        ws.append(list(r))
    ws.append([])
    ws.append(["", "", "", "Grand Total", 35975.00])
    ws.append(["Payment Terms", "NET30"])
    ws.append(["Delivery", "Next-day mandi dispatch"])
    wb.save(path)
    return path


def docx_kirana():
    """Kirana King — same-day delivery, higher prices."""
    from docx import Document

    path = os.path.join(OUT, "kirana_quote.docx")
    doc = Document()
    doc.add_heading("Kirana King Wholesale — Quotation", level=1)
    doc.add_paragraph("To: QuickMart Dark Store, Procurement")
    doc.add_paragraph("Ref: Weekly Dairy & Produce Replenishment - W23")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Item", "Qty", "Unit", "Unit Price (INR)", "Total (INR)"]):
        hdr[i].text = h
    for item, qty, unit, price, total in _rows([19.00, 23.00, 54.00, 310.00]):
        c = table.add_row().cells
        c[0].text, c[1].text, c[2].text = item, str(qty), unit
        c[3].text, c[4].text = f"{price:.2f}", f"{total:.2f}"
    doc.add_paragraph("")
    doc.add_paragraph("Grand Total: INR 39,750.00")
    doc.add_paragraph("Payment Terms: NET30")
    doc.add_paragraph("Delivery: Same-day for dairy. Paneer from our own unit.")
    doc.save(path)
    return path


def pdf_metro():
    """Metro FMCG Supply — digital PDF with a text layer."""
    import pymupdf

    path = os.path.join(OUT, "metro_quote.pdf")
    doc = pymupdf.open()
    page = doc.new_page()
    y = 72
    page.insert_text((72, y), "METRO FMCG SUPPLY", fontsize=16)
    y += 26
    page.insert_text((72, y), "Quotation — RFx W23 (Dairy & Produce)", fontsize=12)
    y += 22
    page.insert_text((72, y), "Date: 2026-05-27   Valid Until: 2026-06-10")
    y += 28
    page.insert_text((72, y), f"{'Item':<20}{'Qty':>5} {'Unit':<5}{'Price':>10}{'Total':>12}")
    y += 16
    page.insert_text((72, y), "-" * 60)
    y += 18
    for item, qty, unit, price, total in _rows([17.50, 21.00, 53.50, 300.00]):
        page.insert_text(
            (72, y), f"{item:<20}{qty:>5} {unit:<5}{price:>10.2f}{total:>12.2f}"
        )
        y += 18
    y += 10
    page.insert_text((72, y), "Grand Total: INR 38,425.00")
    y += 22
    page.insert_text((72, y), "Payment Terms: NET30")
    y += 18
    page.insert_text((72, y), "Delivery: Warehouse pickup, 24 hours")
    doc.save(path)
    doc.close()
    return path


def txt_dailyneeds():
    """Daily Needs Distributors — plain-text email body."""
    path = os.path.join(OUT, "dailyneeds_email.txt")
    content = """Dear Procurement Team,

Thank you for the RFx. Our quotation for the requested items:

Item                Qty    Unit    Unit Price (INR)    Total (INR)
-------------------------------------------------------------------
Tomato              200    kg            18.00            3,600.00
Onion               150    kg            21.50            3,225.00
Full Cream Milk     300    ltr           56.00           16,800.00
Paneer               50    kg           315.00           15,750.00
-------------------------------------------------------------------
Grand Total:                                             39,375.00

Payment Terms: NET15
Delivery: Daily before 6 AM
Validity: 7 days

Regards,
Daily Needs Distributors
+91-98765-43210
"""
    with open(path, "w") as f:
        f.write(content)
    return path


def html_annapurna():
    """Annapurna Foods — HTML email with a pricing table."""
    path = os.path.join(OUT, "annapurna_quote.html")
    rows = "".join(
        f"<tr><td>{i}</td><td>{q}</td><td>{u}</td><td>{p:.2f}</td><td>{t:,.2f}</td></tr>"
        for i, q, u, p, t in _rows([17.00, 22.50, 54.50, 298.00])
    )
    content = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Quotation - Annapurna Foods</title></head>
<body style="font-family:Arial,sans-serif;max-width:640px;margin:auto;">
  <h2 style="color:#c0392b;">Annapurna Foods Pvt Ltd</h2>
  <p>Date: 27 May 2026 &nbsp;|&nbsp; Ref: RFx W23</p>
  <p>Dear Sir/Madam, please find our competitive prices below:</p>
  <table border="1" cellpadding="6" cellspacing="0" style="border-collapse:collapse;width:100%;">
    <thead style="background:#f0f0f0;"><tr>
      <th>Item</th><th>Qty</th><th>Unit</th><th>Unit Price (INR)</th><th>Total (INR)</th>
    </tr></thead>
    <tbody>{rows}</tbody>
  </table>
  <p><strong>Payment Terms:</strong> NET30 &nbsp; <strong>Delivery:</strong> 2 business days</p>
  <p><strong>GST:</strong> 5% applicable</p>
  <p>Best regards,<br>Annapurna Foods Pvt Ltd</p>
</body></html>
"""
    with open(path, "w") as f:
        f.write(content)
    return path


def _price_image(prices, title, color):
    """Render a quotation table to a PIL image (shared by PNG/JPEG/WEBP)."""
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (640, 420), (255, 255, 255))
    d = ImageDraw.Draw(img)
    try:
        big = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 22)
        norm = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 15)
    except OSError:
        big = norm = ImageFont.load_default()
    y = 22
    d.text((24, y), title, fill=color, font=big)
    y += 38
    d.text((24, y), "Quotation — RFx W23", fill=(0, 0, 0), font=norm)
    y += 28
    d.text(
        (24, y),
        f"{'Item':<18}{'Qty':>5} {'Unit':<5}{'Price':>9}{'Total':>11}",
        fill=(0, 0, 0),
        font=norm,
    )
    d.line([(24, y + 20), (616, y + 20)], fill=(0, 0, 0), width=1)
    y += 30
    for item, qty, unit, price, total in _rows(prices):
        d.text(
            (24, y),
            f"{item:<18}{qty:>5} {unit:<5}{price:>9.2f}{total:>11.2f}",
            fill=(0, 0, 0),
            font=norm,
        )
        y += 24
    y += 8
    d.line([(24, y), (616, y)], fill=(0, 0, 0), width=1)
    y += 12
    d.text((24, y), "Payment: NET15  |  Delivery: Next day", fill=(0, 0, 100), font=norm)
    return img


def png_greenvalley():
    """Green Valley — photographed/printed price list (PNG)."""
    path = os.path.join(OUT, "greenvalley_pricelist.png")
    _price_image([15.50, 19.00, 57.00, 312.00], "Green Valley Produce", (0, 110, 0)).save(path)
    return path


def jpg_handwritten():
    """A 'scanned' proforma photo (JPEG) — same items, image-only."""
    path = os.path.join(OUT, "scanned_proforma.jpg")
    _price_image([16.00, 20.50, 55.50, 305.00], "Mandi Direct (Scanned Proforma)", (60, 60, 60)).save(
        path, "JPEG", quality=88
    )
    return path


def webp_vendor():
    """A WEBP price-list photo to cover the webp path."""
    path = os.path.join(OUT, "vendor_pricelist.webp")
    _price_image([17.25, 21.75, 53.00, 308.00], "City Wholesale", (120, 40, 160)).save(path, "WEBP")
    return path


GENERATORS = [
    csv_freshfarm,
    xlsx_sabzimandi,
    docx_kirana,
    pdf_metro,
    txt_dailyneeds,
    html_annapurna,
    png_greenvalley,
    jpg_handwritten,
    webp_vendor,
]


if __name__ == "__main__":
    for gen in GENERATORS:
        p = gen()
        print(f"  wrote {os.path.relpath(p)}  ({os.path.getsize(p)} bytes)")
    print(f"\n{len(GENERATORS)} demo attachments in {OUT}")
