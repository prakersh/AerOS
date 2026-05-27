"""One-time script to generate vendor quote fixture files for integration tests.

Run with: python tests/fixtures/vendor_quotes/generate_fixtures.py
"""

import csv
import os

FIXTURES_DIR = os.path.dirname(os.path.abspath(__file__))


def generate_csv() -> None:
    """Create quote_freshfarm.csv with realistic vendor quote data."""
    path = os.path.join(FIXTURES_DIR, "quote_freshfarm.csv")
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Item", "Qty", "Unit", "Unit Price", "Total"])
        writer.writerow(["Full Cream Milk", 200, "ltr", 56.00, 11200.00])
        writer.writerow(["Paneer (Fresh)", 50, "kg", 320.00, 16000.00])
        writer.writerow(["Curd (Set)", 100, "kg", 45.00, 4500.00])
        writer.writerow(["Butter (Unsalted)", 30, "kg", 480.00, 14400.00])
    print(f"  Created: {path}")


def generate_xlsx() -> None:
    """Create quote_metro.xlsx with a Quotation sheet."""
    from openpyxl import Workbook

    path = os.path.join(FIXTURES_DIR, "quote_metro.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Quotation"
    ws.append(["Item", "Qty", "Unit", "Unit Price (INR)", "Total (INR)"])
    ws.append(["Basmati Rice (Premium)", 100, "kg", 95.00, 9500.00])
    ws.append(["Toor Dal", 50, "kg", 135.00, 6750.00])
    ws.append(["Sunflower Oil", 40, "ltr", 145.00, 5800.00])
    ws.append(["Wheat Flour (Chakki)", 80, "kg", 38.00, 3040.00])
    ws.append(["Sugar (Refined)", 60, "kg", 42.00, 2520.00])

    # Add a summary row
    ws.append([])
    ws.append(["", "", "", "Grand Total", 27610.00])

    wb.save(path)
    print(f"  Created: {path}")


def generate_pdf() -> None:
    """Create quote_bakery.pdf with text-layer pricing data via pymupdf."""
    import pymupdf

    path = os.path.join(FIXTURES_DIR, "quote_bakery.pdf")
    doc = pymupdf.open()
    page = doc.new_page()

    y = 72
    page.insert_text((72, y), "SUNRISE BAKERY SUPPLIES", fontsize=16)
    y += 30
    page.insert_text((72, y), "Quotation for Daily Bakery Items", fontsize=12)
    y += 25
    page.insert_text((72, y), "Date: 2025-05-20    Valid Until: 2025-06-20")
    y += 30

    # Table header
    page.insert_text((72, y), "Item                    Qty    Unit    Unit Price    Total")
    y += 20
    page.insert_text((72, y), "-" * 70)
    y += 18

    items = [
        ("Maida (Refined Flour)", "100", "kg", "34.00", "3400.00"),
        ("Bread Improver", "10", "kg", "280.00", "2800.00"),
        ("Yeast (Instant)", "5", "kg", "450.00", "2250.00"),
        ("Vanilla Essence", "20", "ltr", "120.00", "2400.00"),
    ]
    for item_name, qty, unit, price, total in items:
        line = f"{item_name:<24}{qty:>5}  {unit:<6}  {price:>10}  {total:>10}"
        page.insert_text((72, y), line)
        y += 18

    y += 10
    page.insert_text((72, y), "-" * 70)
    y += 18
    page.insert_text((72, y), f"{'Grand Total':>52}  {'10850.00':>10}")
    y += 30
    page.insert_text((72, y), "Payment Terms: NET30")
    y += 18
    page.insert_text((72, y), "Delivery: Within 48 hours of order confirmation")

    doc.save(path)
    doc.close()
    print(f"  Created: {path}")


def generate_docx() -> None:
    """Create quote_kirana.docx with a pricing table."""
    from docx import Document

    path = os.path.join(FIXTURES_DIR, "quote_kirana.docx")
    doc = Document()
    doc.add_heading("Kirana Store Wholesale - Quotation", level=1)
    doc.add_paragraph("To: AerChain Procurement Team")
    doc.add_paragraph("Date: 20 May 2025")
    doc.add_paragraph("Subject: Weekly Grocery Supply Quote")
    doc.add_paragraph("")

    table = doc.add_table(rows=6, cols=5)
    table.style = "Table Grid"
    headers = ["Item", "Qty", "Unit", "Unit Price (INR)", "Total (INR)"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h

    rows_data = [
        ("Red Onion", "200", "kg", "28.00", "5600.00"),
        ("Potato", "150", "kg", "22.00", "3300.00"),
        ("Tomato", "100", "kg", "35.00", "3500.00"),
        ("Green Chilli", "20", "kg", "65.00", "1300.00"),
        ("Ginger (Fresh)", "15", "kg", "110.00", "1650.00"),
    ]
    for row_idx, (item, qty, unit, price, total) in enumerate(rows_data, start=1):
        table.cell(row_idx, 0).text = item
        table.cell(row_idx, 1).text = qty
        table.cell(row_idx, 2).text = unit
        table.cell(row_idx, 3).text = price
        table.cell(row_idx, 4).text = total

    doc.add_paragraph("")
    doc.add_paragraph("Grand Total: INR 15,350.00")
    doc.add_paragraph("Payment Terms: Cash on Delivery")
    doc.add_paragraph("Delivery: Same day for orders before 10 AM")

    doc.save(path)
    print(f"  Created: {path}")


def generate_png() -> None:
    """Create quote_greenvalley.png with text showing prices using Pillow."""
    from PIL import Image, ImageDraw, ImageFont

    path = os.path.join(FIXTURES_DIR, "quote_greenvalley.png")
    img = Image.new("RGB", (600, 400), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Use default font (available everywhere)
    try:
        font_large = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 20)
        font_normal = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 14)
    except OSError:
        font_large = ImageFont.load_default()
        font_normal = ImageFont.load_default()

    y = 20
    draw.text((20, y), "Green Valley Organic Farms", fill=(0, 100, 0), font=font_large)
    y += 35
    draw.text((20, y), "Quotation - Organic Produce", fill=(0, 0, 0), font=font_normal)
    y += 30

    # Table
    header = f"{'Item':<25} {'Qty':>5} {'Unit':<6} {'Price':>8} {'Total':>10}"
    draw.text((20, y), header, fill=(0, 0, 0), font=font_normal)
    y += 5
    draw.line([(20, y + 15), (580, y + 15)], fill=(0, 0, 0), width=1)
    y += 25

    items = [
        ("Organic Spinach", "50", "kg", "60.00", "3000.00"),
        ("Organic Tomato", "80", "kg", "45.00", "3600.00"),
        ("Organic Cucumber", "40", "kg", "35.00", "1400.00"),
        ("Fresh Coriander", "30", "bunch", "12.00", "360.00"),
    ]
    for item_name, qty, unit, price, total in items:
        line = f"{item_name:<25} {qty:>5} {unit:<6} {price:>8} {total:>10}"
        draw.text((20, y), line, fill=(0, 0, 0), font=font_normal)
        y += 22

    y += 10
    draw.line([(20, y), (580, y)], fill=(0, 0, 0), width=1)
    y += 10
    draw.text((20, y), "Grand Total: INR 8,360.00", fill=(0, 0, 100), font=font_normal)
    y += 25
    draw.text((20, y), "Payment: NET15 | Delivery: Next day", fill=(0, 0, 0), font=font_normal)

    img.save(path)
    print(f"  Created: {path}")


def generate_txt() -> None:
    """Create quote_daily.txt — plain text email body with pricing."""
    path = os.path.join(FIXTURES_DIR, "quote_daily.txt")
    content = """Dear Procurement Team,

Thank you for your enquiry. Please find our quotation for daily essentials below:

Item                    Qty     Unit    Unit Price (INR)    Total (INR)
------------------------------------------------------------------------
Amul Butter (500g)       50     pcs          265.00         13,250.00
Britannia Bread          100    pcs           40.00          4,000.00
Eggs (Farm Fresh)        200    dozen         72.00         14,400.00
Milk (Toned, Pouch)      300    ltr           28.00          8,400.00
------------------------------------------------------------------------
Grand Total:                                                40,050.00

Payment Terms: NET15
Delivery: Daily before 6 AM
Validity: 7 days from date of this quote

Regards,
Daily Needs Distributors
Contact: +91-98765-43210
"""
    with open(path, "w") as f:
        f.write(content)
    print(f"  Created: {path}")


def generate_html() -> None:
    """Create quote_annapurna.html — HTML email with a pricing table."""
    path = os.path.join(FIXTURES_DIR, "quote_annapurna.html")
    content = """<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>Quotation - Annapurna Foods</title></head>
<body style="font-family: Arial, sans-serif; max-width: 600px; margin: auto;">
  <h2 style="color: #c0392b;">Annapurna Foods Pvt Ltd</h2>
  <p>Date: 20 May 2025</p>
  <p>Dear Sir/Madam,</p>
  <p>Please find below our competitive prices for the requested items:</p>

  <table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%;">
    <thead style="background-color: #f0f0f0;">
      <tr>
        <th>Item</th>
        <th>Qty</th>
        <th>Unit</th>
        <th>Unit Price (INR)</th>
        <th>Total (INR)</th>
      </tr>
    </thead>
    <tbody>
      <tr><td>Basmati Rice (1121)</td><td>200</td><td>kg</td><td>88.00</td><td>17,600.00</td></tr>
      <tr><td>Moong Dal (Whole)</td><td>50</td><td>kg</td><td>125.00</td><td>6,250.00</td></tr>
      <tr><td>Turmeric Powder</td><td>25</td><td>kg</td><td>180.00</td><td>4,500.00</td></tr>
      <tr><td>Red Chilli Powder</td><td>20</td><td>kg</td><td>220.00</td><td>4,400.00</td></tr>
      <tr><td>Cumin Seeds</td><td>15</td><td>kg</td><td>340.00</td><td>5,100.00</td></tr>
    </tbody>
    <tfoot>
      <tr style="font-weight: bold;">
        <td colspan="4" style="text-align: right;">Grand Total</td>
        <td>37,850.00</td>
      </tr>
    </tfoot>
  </table>

  <p><strong>Payment Terms:</strong> NET30</p>
  <p><strong>Delivery:</strong> Within 3 business days</p>
  <p><strong>GST:</strong> 5% applicable on all items</p>

  <p>Looking forward to your order.</p>
  <p>Best regards,<br>Annapurna Foods Pvt Ltd</p>
</body>
</html>
"""
    with open(path, "w") as f:
        f.write(content)
    print(f"  Created: {path}")


if __name__ == "__main__":
    print("Generating vendor quote fixtures...")
    generate_csv()
    generate_xlsx()
    generate_pdf()
    generate_docx()
    generate_png()
    generate_txt()
    generate_html()
    print("Done!")
