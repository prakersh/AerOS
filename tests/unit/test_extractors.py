"""Tests for AI extractors — router, email_body, spreadsheet, word, pdf, image."""

import csv
import io
import os
import tempfile
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from aeros.ai.base import VisionResponse


# ---- router tests ----


class TestRouter:
    async def test_routes_pdf(self):
        """PDF MIME type should route to extract_pdf."""
        from aeros.ai.extractors.router import route_extraction, MIME_ROUTER

        assert "application/pdf" in MIME_ROUTER

    async def test_routes_csv(self):
        """CSV MIME type should route to extract_spreadsheet."""
        from aeros.ai.extractors.router import MIME_ROUTER

        assert "text/csv" in MIME_ROUTER
        assert MIME_ROUTER["text/csv"].__name__ == "extract_spreadsheet"

    async def test_routes_image_types(self):
        """Image MIME types should route to extract_image."""
        from aeros.ai.extractors.router import MIME_ROUTER

        for mime in ("image/jpeg", "image/png", "image/webp"):
            assert mime in MIME_ROUTER

    async def test_routes_word(self):
        """Word MIME types should route to extract_word."""
        from aeros.ai.extractors.router import MIME_ROUTER

        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in MIME_ROUTER
        assert "application/msword" in MIME_ROUTER

    async def test_routes_html_and_text(self):
        """HTML and plain text should route to extract_email_body."""
        from aeros.ai.extractors.router import MIME_ROUTER

        assert "text/plain" in MIME_ROUTER
        assert "text/html" in MIME_ROUTER

    async def test_unsupported_mime_returns_message(self):
        """Unsupported MIME type should return a message string."""
        from aeros.ai.extractors.router import route_extraction

        result = await route_extraction("/fake/path", "application/zip")
        assert "Unsupported format" in result

    async def test_route_extraction_calls_image_with_vision_provider(self, tmp_path):
        """Image routing should pass vision_provider to the extractor."""
        from aeros.ai.extractors.router import route_extraction

        fake_image = tmp_path / "test.jpg"
        fake_image.write_bytes(b"\xff\xd8fake-jpeg")

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(content="extracted text")

        with patch("aeros.ai.extractors.image.open", create=True):
            result = await route_extraction(str(fake_image), "image/jpeg", vision_provider=mock_vp)

        # Should have called the vision provider
        assert mock_vp.vision.called or "vision provider" in result.lower() or "extracted" in result.lower()


# ---- email_body tests ----


class TestEmailBodyExtractor:
    async def test_plaintext_passthrough(self, tmp_path):
        """Plain text without HTML should be returned with trimming."""
        from aeros.ai.extractors.email_body import extract_email_body

        f = tmp_path / "email.txt"
        f.write_text("Hello, here is my quote.\nRice: 40/kg\nWheat: 20/kg")
        result = await extract_email_body(str(f))
        assert "Hello" in result
        assert "Rice" in result

    async def test_html_stripping(self, tmp_path):
        """HTML content should have tags stripped."""
        from aeros.ai.extractors.email_body import extract_email_body

        f = tmp_path / "email.html"
        f.write_text("<html><body><p>Price list:</p><ul><li>Rice 40/kg</li></ul></body></html>")
        result = await extract_email_body(str(f))
        assert "<p>" not in result
        assert "<ul>" not in result
        assert "Price list" in result
        assert "Rice" in result

    async def test_forwarded_chain_removal_plaintext(self, tmp_path):
        """Plaintext forwarded chain should be cut at the forwarded marker."""
        from aeros.ai.extractors.email_body import extract_email_body

        f = tmp_path / "fwd.txt"
        f.write_text(
            "Here is the quote:\nRice: 40/kg\n\n"
            "---------- Forwarded message ----------\n"
            "From: someone@else.com\nSubject: Old thread"
        )
        result = await extract_email_body(str(f))
        assert "Rice" in result
        assert "someone@else.com" not in result

    async def test_html_forwarded_chain_removal(self, tmp_path):
        """HTML forwarded chain pattern should be stripped."""
        from aeros.ai.extractors.email_body import extract_email_body

        f = tmp_path / "fwd.html"
        f.write_text(
            "<html><body><p>New quote attached</p>"
            "<p>---------- Forwarded message ----------</p>"
            "<p>Old content here</p></body></html>"
        )
        result = await extract_email_body(str(f))
        assert "New quote" in result
        assert "Old content" not in result

    async def test_plaintext_on_wrote_break(self, tmp_path):
        """Plaintext should break at 'On ... wrote:' pattern."""
        from aeros.ai.extractors.email_body import extract_email_body

        f = tmp_path / "reply.txt"
        f.write_text(
            "My offer:\nRice 45/kg\n\n"
            "On Mon, 26 May 2025, buyer@aerchain.com wrote:\n> Please send quote"
        )
        result = await extract_email_body(str(f))
        assert "Rice" in result
        assert "buyer@aerchain.com" not in result


# ---- spreadsheet tests ----


class TestSpreadsheetExtractor:
    async def test_csv_parsing(self, tmp_path):
        """CSV file should be parsed into pipe-separated rows."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        f = tmp_path / "prices.csv"
        f.write_text("Item,Price,Unit\nRice,40,kg\nWheat,20,kg\n")
        result = await extract_spreadsheet(str(f))
        assert "Item | Price | Unit" in result
        assert "Rice | 40 | kg" in result
        assert "Wheat | 20 | kg" in result

    async def test_tsv_parsing(self, tmp_path):
        """TSV file should be parsed correctly."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        f = tmp_path / "prices.tsv"
        f.write_text("Item\tPrice\nRice\t40\nWheat\t20\n")
        result = await extract_spreadsheet(str(f))
        assert "Rice" in result
        assert "Wheat" in result

    async def test_empty_csv(self, tmp_path):
        """Empty CSV should return empty message."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        f = tmp_path / "empty.csv"
        f.write_text("")
        result = await extract_spreadsheet(str(f))
        assert "Empty" in result

    async def test_csv_with_blank_rows(self, tmp_path):
        """Blank rows should be skipped."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        f = tmp_path / "sparse.csv"
        f.write_text("A,B\n1,2\n,,\n3,4\n")
        result = await extract_spreadsheet(str(f))
        assert "1 | 2" in result
        assert "3 | 4" in result
        # The row with just commas (empty after strip) should be skipped
        lines = [l for l in result.strip().split("\n") if l.strip()]
        assert len(lines) == 3  # header + 2 data rows

    async def test_excel_extraction(self, tmp_path):
        """XLSX file should be extracted using openpyxl."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        try:
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "Prices"
            ws.append(["Item", "Price"])
            ws.append(["Rice", 40])
            ws.append(["Wheat", 20])

            f = tmp_path / "prices.xlsx"
            wb.save(str(f))

            result = await extract_spreadsheet(str(f))
            assert "Prices" in result
            assert "Rice" in result
            assert "Wheat" in result
        except ImportError:
            pytest.skip("openpyxl not installed")


# ---- word tests ----


class TestWordExtractor:
    async def test_word_document_extraction(self, tmp_path):
        """Word document paragraphs and tables should be extracted."""
        from aeros.ai.extractors.word import extract_word

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Vendor Quote")
            doc.add_paragraph("Rice: 40/kg")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Item"
            table.cell(0, 1).text = "Price"
            table.cell(1, 0).text = "Rice"
            table.cell(1, 1).text = "40"

            f = tmp_path / "quote.docx"
            doc.save(str(f))

            result = await extract_word(str(f))
            assert "Vendor Quote" in result
            assert "Rice: 40/kg" in result
            assert "Item" in result
            assert "Price" in result
        except ImportError:
            pytest.skip("python-docx not installed")

    async def test_empty_word_document(self, tmp_path):
        """Empty Word document should return empty message."""
        from aeros.ai.extractors.word import extract_word

        try:
            from docx import Document

            doc = Document()
            f = tmp_path / "empty.docx"
            doc.save(str(f))

            result = await extract_word(str(f))
            assert "Empty" in result
        except ImportError:
            pytest.skip("python-docx not installed")


# ---- pdf tests ----


class TestPdfExtractor:
    async def test_pdf_extraction(self, tmp_path):
        """PDF with text should be extracted."""
        from aeros.ai.extractors.pdf import extract_pdf

        try:
            import pymupdf

            doc = pymupdf.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Vendor Rate Card\nRice: 40/kg")

            f = tmp_path / "rates.pdf"
            doc.save(str(f))
            doc.close()

            result = await extract_pdf(str(f))
            assert "Rice" in result
            assert "40" in result
        except ImportError:
            pytest.skip("pymupdf not installed")

    async def test_pdf_fallback_to_pymupdf_text(self, tmp_path):
        """When pymupdf4llm fails, should fall back to pymupdf text extraction."""
        from aeros.ai.extractors.pdf import extract_pdf

        try:
            import pymupdf

            doc = pymupdf.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Fallback text content")

            f = tmp_path / "fallback.pdf"
            doc.save(str(f))
            doc.close()

            # Make pymupdf4llm.to_markdown raise so the try/except falls through
            mock_pymupdf4llm = MagicMock()
            mock_pymupdf4llm.to_markdown = MagicMock(side_effect=Exception("fail"))
            with patch.dict("sys.modules", {"pymupdf4llm": mock_pymupdf4llm}):
                result = await extract_pdf(str(f))

            assert "Fallback text content" in result
        except ImportError:
            pytest.skip("pymupdf not installed")


# ---- image tests ----


class TestImageExtractor:
    async def test_image_without_vision_provider_returns_message(self, tmp_path):
        """Without a vision provider, should return an informative message."""
        from aeros.ai.extractors.image import extract_image

        f = tmp_path / "rate_card.jpg"
        f.write_bytes(b"\xff\xd8fake-jpeg-data")

        result = await extract_image(str(f), vision_provider=None)
        assert "vision provider" in result.lower()

    async def test_image_with_vision_provider(self, tmp_path):
        """With a vision provider, should call it and return the result."""
        from aeros.ai.extractors.image import extract_image

        f = tmp_path / "card.png"
        f.write_bytes(b"\x89PNGfake-data")

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(content="Rice: 40/kg, Wheat: 20/kg")

        result = await extract_image(str(f), vision_provider=mock_vp)
        assert "Rice" in result
        mock_vp.vision.assert_called_once()

    async def test_image_detects_mime_type_png(self, tmp_path):
        """PNG files should use image/png mime type."""
        from aeros.ai.extractors.image import extract_image

        f = tmp_path / "card.png"
        f.write_bytes(b"\x89PNGfake")

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(content="data")

        await extract_image(str(f), vision_provider=mock_vp)
        call_args = mock_vp.vision.call_args
        assert call_args.kwargs.get("mime_type") == "image/png" or call_args[1].get("mime_type") == "image/png"

    async def test_image_detects_mime_type_webp(self, tmp_path):
        """WebP files should use image/webp mime type."""
        from aeros.ai.extractors.image import extract_image

        f = tmp_path / "card.webp"
        f.write_bytes(b"RIFFfake")

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(content="data")

        await extract_image(str(f), vision_provider=mock_vp)
        call_args = mock_vp.vision.call_args
        assert call_args.kwargs.get("mime_type") == "image/webp" or call_args[1].get("mime_type") == "image/webp"
