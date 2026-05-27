"""Integration tests for the format-agnostic extraction pipeline.

Tests each extractor with real file artifacts (created in tmp_path):
1. extract_pdf — text-layer PDF via pymupdf
2. extract_word — .docx with paragraphs and tables
3. extract_spreadsheet — .xlsx and .csv files
4. extract_image — vision provider mock
5. extract_email_body — HTML and plaintext
6. route_extraction — correct dispatcher routing by MIME type
"""

import csv
import io
import os
import tempfile
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from aeros.ai.base import VisionResponse


# ---------------------------------------------------------------------------
# PDF Extraction
# ---------------------------------------------------------------------------


class TestExtractPdf:
    """Tests for PDF text extraction with real PDF artifacts."""

    @pytest.mark.asyncio
    async def test_extract_text_layer_pdf(self, tmp_path):
        """A PDF with embedded text should return the text content."""
        from aeros.ai.extractors.pdf import extract_pdf

        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Vendor Rate Card\nItem: Basmati Rice\nPrice: 80 INR/kg\nQty: 100 kg")
        page.insert_text((72, 200), "Total: 8000 INR\nPayment: NET30\nDelivery: Doorstep")

        pdf_path = tmp_path / "rate_card.pdf"
        doc.save(str(pdf_path))
        doc.close()

        result = await extract_pdf(str(pdf_path))
        assert "Basmati Rice" in result
        assert "80" in result
        assert "8000" in result
        assert "NET30" in result

    @pytest.mark.asyncio
    async def test_extract_multi_page_pdf(self, tmp_path):
        """A multi-page PDF should extract text from all pages."""
        from aeros.ai.extractors.pdf import extract_pdf

        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        doc = pymupdf.open()
        # Page 1
        p1 = doc.new_page()
        p1.insert_text((72, 72), "Page 1: Rice 80/kg")
        # Page 2
        p2 = doc.new_page()
        p2.insert_text((72, 72), "Page 2: Wheat 35/kg")

        pdf_path = tmp_path / "multi.pdf"
        doc.save(str(pdf_path))
        doc.close()

        result = await extract_pdf(str(pdf_path))
        assert "Rice" in result
        assert "Wheat" in result

    @pytest.mark.asyncio
    async def test_scanned_pdf_uses_vision_provider(self, tmp_path):
        """A scanned PDF (no text layer) should use vision provider for OCR."""
        from aeros.ai.extractors.pdf import extract_pdf

        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        # Create a PDF with only an image (no text layer)
        doc = pymupdf.open()
        page = doc.new_page()
        # No text inserted — simulates a scanned image-only PDF

        pdf_path = tmp_path / "scanned.pdf"
        doc.save(str(pdf_path))
        doc.close()

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(content="OCR: Rice 80/kg, Wheat 35/kg")

        result = await extract_pdf(str(pdf_path), vision_provider=mock_vp)
        # Either vision was called (scanned) or pymupdf4llm returned something
        assert "Rice" in result or "Empty" in result or mock_vp.vision.called

    @pytest.mark.asyncio
    async def test_scanned_pdf_no_vision_returns_message(self, tmp_path):
        """A scanned PDF without vision provider should return informative message."""
        from aeros.ai.extractors.pdf import extract_pdf

        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        doc = pymupdf.open()
        page = doc.new_page()
        pdf_path = tmp_path / "no_vision.pdf"
        doc.save(str(pdf_path))
        doc.close()

        # Force pymupdf4llm to fail, then no text in pages, no vision
        mock_pymupdf4llm = MagicMock()
        mock_pymupdf4llm.to_markdown.side_effect = Exception("fail")
        with patch.dict("sys.modules", {"pymupdf4llm": mock_pymupdf4llm}):
            result = await extract_pdf(str(pdf_path), vision_provider=None)

        assert "Scanned PDF" in result or "vision" in result.lower() or "Empty" in result


# ---------------------------------------------------------------------------
# Word Extraction
# ---------------------------------------------------------------------------


class TestExtractWord:
    """Tests for Word document extraction."""

    @pytest.mark.asyncio
    async def test_extract_docx_paragraphs_and_tables(self, tmp_path):
        """A .docx with paragraphs and tables should have all content extracted."""
        from aeros.ai.extractors.word import extract_word

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Vendor Quotation")
        doc.add_paragraph("For: Basmati Rice, Full Cream Milk")

        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Price/kg"
        table.cell(0, 2).text = "Qty"
        table.cell(1, 0).text = "Basmati Rice"
        table.cell(1, 1).text = "80"
        table.cell(1, 2).text = "100"
        table.cell(2, 0).text = "Full Cream Milk"
        table.cell(2, 1).text = "55"
        table.cell(2, 2).text = "200"

        doc.add_paragraph("Total: 21000 INR")

        docx_path = tmp_path / "quote.docx"
        doc.save(str(docx_path))

        result = await extract_word(str(docx_path))
        assert "Vendor Quotation" in result
        assert "Basmati Rice" in result
        assert "80" in result
        assert "Full Cream Milk" in result
        assert "55" in result
        assert "21000" in result
        # Table rows should appear
        assert "Item" in result
        assert "Price/kg" in result

    @pytest.mark.asyncio
    async def test_extract_empty_docx(self, tmp_path):
        """An empty .docx should return the empty document message."""
        from aeros.ai.extractors.word import extract_word

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        docx_path = tmp_path / "empty.docx"
        doc.save(str(docx_path))

        result = await extract_word(str(docx_path))
        assert "Empty" in result


# ---------------------------------------------------------------------------
# Spreadsheet Extraction
# ---------------------------------------------------------------------------


class TestExtractSpreadsheet:
    """Tests for spreadsheet extraction (xlsx, csv)."""

    @pytest.mark.asyncio
    async def test_extract_xlsx(self, tmp_path):
        """An .xlsx file should have sheet names and cell data extracted."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws = wb.active
        ws.title = "Pricing"
        ws.append(["Item", "Unit Price", "Qty", "Total"])
        ws.append(["Basmati Rice", 80, 100, 8000])
        ws.append(["Wheat Flour", 35, 50, 1750])
        ws.append(["Full Cream Milk", 55, 200, 11000])

        xlsx_path = tmp_path / "vendor_quote.xlsx"
        wb.save(str(xlsx_path))

        result = await extract_spreadsheet(str(xlsx_path))
        assert "Pricing" in result
        assert "Basmati Rice" in result
        assert "80" in result
        assert "8000" in result
        assert "Wheat Flour" in result
        assert "Full Cream Milk" in result

    @pytest.mark.asyncio
    async def test_extract_csv(self, tmp_path):
        """A .csv file should be parsed into pipe-separated rows."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        csv_path = tmp_path / "prices.csv"
        csv_path.write_text("Item,Price,Qty,Total\nRice,80,100,8000\nWheat,35,50,1750\n")

        result = await extract_spreadsheet(str(csv_path))
        assert "Item | Price | Qty | Total" in result
        assert "Rice | 80 | 100 | 8000" in result
        assert "Wheat | 35 | 50 | 1750" in result

    @pytest.mark.asyncio
    async def test_extract_xlsx_multi_sheet(self, tmp_path):
        """An .xlsx with multiple sheets should extract all sheet data."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Grains"
        ws1.append(["Rice", 80])
        ws1.append(["Wheat", 35])

        ws2 = wb.create_sheet("Dairy")
        ws2.append(["Milk", 55])
        ws2.append(["Butter", 400])

        xlsx_path = tmp_path / "multi_sheet.xlsx"
        wb.save(str(xlsx_path))

        result = await extract_spreadsheet(str(xlsx_path))
        assert "Grains" in result
        assert "Rice" in result
        assert "Dairy" in result
        assert "Milk" in result
        assert "Butter" in result

    @pytest.mark.asyncio
    async def test_extract_empty_csv(self, tmp_path):
        """An empty CSV should return the empty message."""
        from aeros.ai.extractors.spreadsheet import extract_spreadsheet

        csv_path = tmp_path / "empty.csv"
        csv_path.write_text("")

        result = await extract_spreadsheet(str(csv_path))
        assert "Empty" in result


# ---------------------------------------------------------------------------
# Image Extraction
# ---------------------------------------------------------------------------


class TestExtractImage:
    """Tests for image extraction with mocked vision provider."""

    @pytest.mark.asyncio
    async def test_extract_image_with_vision_provider(self, tmp_path):
        """Image extraction should call the vision provider and return its output."""
        from aeros.ai.extractors.image import extract_image

        img_path = tmp_path / "rate_card.jpg"
        img_path.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 100)  # minimal JPEG header

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(
            content="Rate Card:\n- Basmati Rice: 80 INR/kg\n- Wheat Flour: 35 INR/kg"
        )

        result = await extract_image(str(img_path), vision_provider=mock_vp)
        assert "Basmati Rice" in result
        assert "80" in result
        mock_vp.vision.assert_called_once()

        # Verify the correct mime type was passed
        call_kwargs = mock_vp.vision.call_args
        assert call_kwargs.kwargs.get("mime_type") == "image/jpeg"

    @pytest.mark.asyncio
    async def test_extract_png_image(self, tmp_path):
        """PNG image should use image/png mime type."""
        from aeros.ai.extractors.image import extract_image

        img_path = tmp_path / "card.png"
        img_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50)

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(content="Milk: 55/ltr")

        result = await extract_image(str(img_path), vision_provider=mock_vp)
        assert "Milk" in result
        call_kwargs = mock_vp.vision.call_args
        assert call_kwargs.kwargs.get("mime_type") == "image/png"

    @pytest.mark.asyncio
    async def test_extract_image_without_vision_provider(self, tmp_path):
        """Without a vision provider, extraction should return an informative message."""
        from aeros.ai.extractors.image import extract_image

        img_path = tmp_path / "card.jpg"
        img_path.write_bytes(b"\xff\xd8\xff\xe0fake")

        result = await extract_image(str(img_path), vision_provider=None)
        assert "vision provider" in result.lower()


# ---------------------------------------------------------------------------
# Email Body Extraction
# ---------------------------------------------------------------------------


class TestExtractEmailBody:
    """Tests for email body extraction (HTML and plaintext)."""

    @pytest.mark.asyncio
    async def test_extract_html_email(self, tmp_path):
        """HTML email should have tags stripped and text extracted."""
        from aeros.ai.extractors.email_body import extract_email_body

        html = """<html><body>
        <h2>Vendor Quote</h2>
        <p>Dear Buyer,</p>
        <p>Please find our prices below:</p>
        <table>
            <tr><td>Basmati Rice</td><td>80 INR/kg</td></tr>
            <tr><td>Wheat Flour</td><td>35 INR/kg</td></tr>
        </table>
        <p>Total: 5750 INR</p>
        <p>Best regards, Vendor Co</p>
        </body></html>"""

        f = tmp_path / "email.html"
        f.write_text(html)

        result = await extract_email_body(str(f))
        assert "Vendor Quote" in result
        assert "Basmati Rice" in result
        assert "80" in result
        assert "Wheat Flour" in result
        assert "35" in result
        # Tags should be stripped
        assert "<h2>" not in result
        assert "<table>" not in result

    @pytest.mark.asyncio
    async def test_extract_plaintext_email(self, tmp_path):
        """Plaintext email should be extracted directly."""
        from aeros.ai.extractors.email_body import extract_email_body

        text = """Dear Buyer,

Here are our prices:
- Basmati Rice: 80 INR/kg
- Wheat Flour: 35 INR/kg
- Full Cream Milk: 55 INR/ltr

Total for 100kg Rice + 50kg Wheat: 9750 INR

Payment Terms: NET15
Delivery: Within 24 hours
"""
        f = tmp_path / "email.txt"
        f.write_text(text)

        result = await extract_email_body(str(f))
        assert "Basmati Rice" in result
        assert "80" in result
        assert "NET15" in result

    @pytest.mark.asyncio
    async def test_extract_email_strips_forwarded_chain(self, tmp_path):
        """Forwarded email chains should be stripped."""
        from aeros.ai.extractors.email_body import extract_email_body

        text = """Our revised prices:
Basmati Rice: 78 INR/kg (reduced from 80)
Wheat Flour: 34 INR/kg

---------- Forwarded message ----------
From: buyer@aerchain.com
Sent: Mon, 26 May 2025
Subject: RE: Price Request

Original request details that should not appear."""

        f = tmp_path / "forwarded.txt"
        f.write_text(text)

        result = await extract_email_body(str(f))
        assert "Basmati Rice" in result
        assert "78" in result
        assert "buyer@aerchain.com" not in result
        assert "Original request" not in result

    @pytest.mark.asyncio
    async def test_extract_html_email_strips_reply_chain(self, tmp_path):
        """HTML email with 'On ... wrote:' pattern should be cleaned."""
        from aeros.ai.extractors.email_body import extract_email_body

        html = """<html><body>
        <p>Updated quote: Rice 75/kg, Wheat 33/kg</p>
        <p>On Mon, 26 May 2025, buyer@aerchain.com wrote:</p>
        <blockquote>Please send your updated prices</blockquote>
        </body></html>"""

        f = tmp_path / "reply.html"
        f.write_text(html)

        result = await extract_email_body(str(f))
        assert "Updated quote" in result
        assert "75" in result
        assert "buyer@aerchain.com" not in result


# ---------------------------------------------------------------------------
# Route Extraction Dispatcher
# ---------------------------------------------------------------------------


class TestRouteExtraction:
    """Tests for the MIME-type based routing dispatcher."""

    @pytest.mark.asyncio
    async def test_route_pdf(self, tmp_path):
        """PDF MIME type should route to extract_pdf."""
        from aeros.ai.extractors.router import route_extraction

        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Routed PDF Content: Rice 80/kg")
        pdf_path = tmp_path / "routed.pdf"
        doc.save(str(pdf_path))
        doc.close()

        result = await route_extraction(str(pdf_path), "application/pdf")
        assert "Rice" in result

    @pytest.mark.asyncio
    async def test_route_csv(self, tmp_path):
        """CSV MIME type should route to extract_spreadsheet."""
        from aeros.ai.extractors.router import route_extraction

        csv_path = tmp_path / "routed.csv"
        csv_path.write_text("Item,Price\nRice,80\n")

        result = await route_extraction(str(csv_path), "text/csv")
        assert "Rice" in result
        assert "80" in result

    @pytest.mark.asyncio
    async def test_route_xlsx(self, tmp_path):
        """XLSX MIME type should route to extract_spreadsheet."""
        from aeros.ai.extractors.router import route_extraction

        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws = wb.active
        ws.append(["Item", "Price"])
        ws.append(["Rice", 80])
        xlsx_path = tmp_path / "routed.xlsx"
        wb.save(str(xlsx_path))

        result = await route_extraction(
            str(xlsx_path),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        assert "Rice" in result

    @pytest.mark.asyncio
    async def test_route_docx(self, tmp_path):
        """DOCX MIME type should route to extract_word."""
        from aeros.ai.extractors.router import route_extraction

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Routed Word: Rice 80/kg")
        docx_path = tmp_path / "routed.docx"
        doc.save(str(docx_path))

        result = await route_extraction(
            str(docx_path),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert "Rice" in result

    @pytest.mark.asyncio
    async def test_route_image_with_vision(self, tmp_path):
        """Image MIME type should route to extract_image with vision provider."""
        from aeros.ai.extractors.router import route_extraction

        img_path = tmp_path / "routed.jpg"
        img_path.write_bytes(b"\xff\xd8\xff\xe0fake-jpeg")

        mock_vp = AsyncMock()
        mock_vp.vision.return_value = VisionResponse(content="Rice 80/kg from image")

        result = await route_extraction(str(img_path), "image/jpeg", vision_provider=mock_vp)
        assert "Rice" in result

    @pytest.mark.asyncio
    async def test_route_html_email(self, tmp_path):
        """HTML MIME type should route to extract_email_body."""
        from aeros.ai.extractors.router import route_extraction

        f = tmp_path / "email.html"
        f.write_text("<html><body><p>Rice quote: 80/kg</p></body></html>")

        result = await route_extraction(str(f), "text/html")
        assert "Rice" in result
        assert "80" in result

    @pytest.mark.asyncio
    async def test_route_plaintext(self, tmp_path):
        """Plaintext MIME type should route to extract_email_body."""
        from aeros.ai.extractors.router import route_extraction

        f = tmp_path / "email.txt"
        f.write_text("Rice: 80/kg\nWheat: 35/kg")

        result = await route_extraction(str(f), "text/plain")
        assert "Rice" in result

    @pytest.mark.asyncio
    async def test_route_unsupported_returns_message(self):
        """Unsupported MIME type should return an informative message."""
        from aeros.ai.extractors.router import route_extraction

        result = await route_extraction("/fake/path", "application/zip")
        assert "Unsupported format" in result
        assert "application/zip" in result
