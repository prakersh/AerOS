"""Word document extractor."""

from zipfile import BadZipFile

from docx import Document


async def extract_word(file_path: str, **kwargs) -> str:
    try:
        doc = Document(file_path)
    except BadZipFile:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            text = f.read()
        return text if text.strip() else "[Unreadable Word document — binary .doc format]"

    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts) if parts else "[Empty Word document]"
